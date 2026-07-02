from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Cathay_Praxis_Law_Firm_Profile_CN_EN.docx"
NAVY = RGBColor(14, 41, 65)
GOLD = RGBColor(174, 139, 71)
GRAY = RGBColor(92, 101, 111)
LIGHT = "EAF0F4"
CH_FONT = "PingFang SC"
EN_FONT = "Aptos"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, bold=False, color=None, font=CH_FONT, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color


def style_para(p, before=0, after=6, line=1.18, keep=False):
    f = p.paragraph_format
    f.space_before = Pt(before); f.space_after = Pt(after); f.line_spacing = line
    f.keep_with_next = keep


def add_text(doc, zh, en=None, size=10.5, after=5, indent=True):
    p = doc.add_paragraph(); style_para(p, after=after, line=1.22)
    if indent: p.paragraph_format.first_line_indent = Inches(0.25)
    set_run(p.add_run(zh), size=size)
    if en:
        p2 = doc.add_paragraph(); style_para(p2, after=after + 3, line=1.15)
        if indent: p2.paragraph_format.first_line_indent = Inches(0.25)
        set_run(p2.add_run(en), size=size - 0.3, color=GRAY, font=EN_FONT)
    return p


def add_heading(doc, zh, en, level=1):
    p = doc.add_paragraph(style=f"Heading {level}"); style_para(p, before=14 if level == 1 else 10, after=6, keep=True)
    set_run(p.add_run(zh), size=16 if level == 1 else 12.5, bold=True, color=NAVY)
    set_run(p.add_run(f"  {en}"), size=10.5 if level == 1 else 9.5, bold=False, color=GOLD, font=EN_FONT)
    return p


def add_bilingual_bullets(doc, items_zh, items_en):
    for zh, en in zip(items_zh, items_en):
        p = doc.add_paragraph(style="List Bullet"); style_para(p, after=3, line=1.12)
        set_run(p.add_run(zh), size=10.1)
        r = p.add_run(f"\n{en}"); set_run(r, size=9.4, color=GRAY, font=EN_FONT)


def add_footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_para(p, after=0)
    set_run(p.add_run("CATHAY PRAXIS LAW, PC  |  博衡律师事务所  |  cathaypraxislaw.com"), size=8, color=GRAY, font=EN_FONT)


def add_profile_header(doc, person):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.65); table.columns[1].width = Inches(4.85)
    left, right = table.rows[0].cells
    left.width = Inches(1.65); right.width = Inches(4.85)
    for c in (left, right):
        set_cell_margins(c, 80, 100, 80, 100); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.paragraphs[0].add_run().add_picture(str(ROOT / person["image"]), width=Inches(1.35), height=Inches(1.55))
    set_cell_shading(right, LIGHT)
    p = right.paragraphs[0]; style_para(p, after=3)
    set_run(p.add_run(person["name_zh"]), 20, True, NAVY)
    set_run(p.add_run(f"  {person['name_en']}"), 13, False, GOLD, EN_FONT)
    p = right.add_paragraph(); style_para(p, after=3)
    set_run(p.add_run(person["role_zh"]), 11, True, NAVY)
    set_run(p.add_run(f"  |  {person['role_en']}"), 10, False, GRAY, EN_FONT)
    p = right.add_paragraph(); style_para(p, after=2)
    set_run(p.add_run(f"{person['location_zh']}  ·  {person['location_en']}"), 9.2, color=GRAY)
    p = right.add_paragraph(); style_para(p, after=0)
    set_run(p.add_run(person["email"]), 9.2, color=GRAY, font=EN_FONT)


people = [
    dict(name_zh="周争平", name_en="Zhengping (David) Zhou", role_zh="创始人、执业律师", role_en="Founder & Attorney", location_zh="中国长沙", location_en="Changsha, China", email="zhengping.zhou@cathaypraxislaw.com", image="images/david-zhou.png",
         bio_zh="周争平律师拥有美国加州及中华人民共和国律师执业资格。近二十年来，他为跨中美及亚洲地区运营的企业、投资者及机构提供复杂法律与监管事务咨询，专注于跨境投资与并购、美国出口管制与经济制裁、监管合规及国际争议解决。他善于将严谨的法律分析与对商业环境的务实理解相结合，为客户提供清晰、可执行并服务于商业目标的法律方案。",
         bio_en="David Zhou is admitted to practice in California and the People’s Republic of China. For nearly two decades, he has advised companies, investors, and institutions operating across the United States and Asia on complex legal and regulatory matters. His practice focuses on cross-border investment and M&A, U.S. export controls and economic sanctions, regulatory compliance, and international dispute resolution. He combines rigorous legal analysis with a practical understanding of commercial realities to deliver clear, executable advice aligned with clients’ business objectives.",
         practices_zh=["跨境投资与并购（M&A）", "美国出口管制、经济制裁与贸易合规", "国际仲裁与跨境争议解决", "公司治理与美国市场准入", "监管合规与风险管理"],
         practices_en=["Cross-Border Investment & M&A", "U.S. Export Controls, Economic Sanctions & Trade Compliance", "International Arbitration & Cross-Border Dispute Resolution", "Corporate Governance & U.S. Market Entry", "Regulatory Compliance & Risk Management"],
         matters_zh=["就医疗及工业领域上市公司的大型跨境交易提供法律服务", "为企业进入欧美市场设计出口管制及制裁合规体系", "参与涉及跨境投资与贸易的国际仲裁案件", "就美国监管执法行动及合规整改提供战略咨询"],
         matters_en=["Advised on large-scale cross-border transactions in the healthcare and industrial sectors", "Designed export-control and sanctions-compliance frameworks for operations in the United States and Europe", "Acted in international arbitration proceedings involving cross-border investment and trade", "Provided strategic advice on U.S. regulatory enforcement and compliance remediation"],
         education_zh=["美国加州大学伯克利分校 法学硕士", "中南大学 法律硕士"], education_en=["LL.M., University of California, Berkeley", "Master of Law, Central South University"],
         bars_zh="美国加州律师执业资格；中华人民共和国律师执业资格", bars_en="State Bar of California; PRC Bar",
         recognition_zh=["《商法》A-List法律精英——国际贸易与合规领域（2023年、2025年）", "在跨境交易与争议解决领域获得多项专业认可"], recognition_en=["China Business Law Journal A-List — International Trade & Compliance (2023 and 2025)", "Recognized for work in cross-border transactions and dispute resolution"]),
    dict(name_zh="叶明博", name_en="Mingbo Ye", role_zh="顾问律师", role_en="Of Counsel", location_zh="美国得克萨斯州休斯敦", location_en="Houston, Texas", email="mingbo.ye@cathaypraxislaw.com", image="images/mingbo-ye.png",
         bio_zh="叶明博律师拥有美国得克萨斯州与加州律师执业资格，专注于知识产权诉讼、商标申请及组合管理。她兼具专利审查与美国诉讼代理的双重背景：曾在中国国家知识产权局担任专利审查员七年，后在美国知名知识产权律所 Kilpatrick Townsend & Stockton LLP 执业三年。叶律师熟悉中美两国知识产权制度，在 Schedule A 执法行动以及专利、商标和版权侵权纠纷领域具有丰富经验。",
         bio_en="Mingbo Ye is admitted in Texas and California. Her practice focuses on intellectual property litigation, trademark prosecution, and portfolio management. She brings a distinctive combination of patent-examination and U.S. litigation experience, having served for seven years as a Patent Examiner at the China National Intellectual Property Administration before practicing for three years at Kilpatrick Townsend & Stockton LLP. She is experienced in Schedule A enforcement actions and patent, trademark, and copyright disputes.",
         practices_zh=["知识产权诉讼", "商标申请与组合管理", "跨境知识产权战略与争议解决"], practices_en=["Intellectual Property Litigation", "Trademark Prosecution & Portfolio Management", "Cross-Border IP Strategy & Dispute Resolution"],
         matters_zh=["在伊利诺伊州北区联邦法院及得州多个联邦地区法院代理知识产权诉讼", "处理涉及专利、商标及版权侵权的 Schedule A 执法行动", "为跨境品牌提供商标组合管理及争议解决支持"], matters_en=["Represented clients in IP litigation in the Northern District of Illinois and multiple federal districts in Texas", "Handled Schedule A enforcement actions involving patent, trademark, and copyright claims", "Supported cross-border brands with trademark portfolio management and dispute resolution"],
         education_zh=["加州大学洛杉矶分校（UCLA）法学院", "北京大学", "华东政法大学"], education_en=["UCLA School of Law", "Peking University", "East China University of Political Science and Law"],
         bars_zh="美国得克萨斯州执业律师；美国加州执业律师；中国专利代理人资格", bars_en="State Bar of Texas; State Bar of California; Qualified Patent Attorney, PRC",
         recognition_zh=["前中国国家知识产权局专利审查员", "曾任职于 Kilpatrick Townsend & Stockton LLP"], recognition_en=["Former Patent Examiner, China National Intellectual Property Administration", "Former attorney at Kilpatrick Townsend & Stockton LLP"]),
    dict(name_zh="周媛", name_en="Yuan Zhou", role_zh="顾问律师", role_en="Of Counsel", location_zh="中国深圳", location_en="Shenzhen, China", email="yuan.zhou@cathaypraxislaw.com", image="images/zhou-yuan.jpg",
         bio_zh="周媛律师拥有中国大陆、美国加州及英格兰与威尔士三重执业资格，曾在顶尖中资律师事务所执业逾五年，专注于高标的额跨境争议解决。她毕业于香港中文大学法律博士项目，曾获国家留学基金委公派资助赴瑞士日内瓦大学深造，并荣获广东省五一劳动奖章。周律师曾主导管理八起单案标的额超亿元人民币的重大争议案件。",
         bio_en="Yuan Zhou is admitted in Mainland China, California, and England & Wales. With more than five years of experience at leading Chinese law firms, she focuses on high-stakes cross-border dispute resolution. She holds a J.D. from The Chinese University of Hong Kong, studied at the University of Geneva as a Chinese government-sponsored scholar, and received the Guangdong Province May 1st Labor Medal. She has led eight disputes, each involving claims exceeding RMB 100 million.",
         practices_zh=["跨境诉讼与国际仲裁", "国际贸易与反倾销调查应对", "建设工程争议", "跨境投资、并购及尽职调查", "境外判决与仲裁裁决的承认及执行"], practices_en=["Cross-Border Litigation & International Arbitration", "International Trade & Anti-Dumping Investigations", "Construction & Engineering Disputes", "Cross-Border Investment, M&A & Due Diligence", "Recognition & Enforcement of Foreign Judgments and Arbitral Awards"],
         matters_zh=["为美国贸易公司制定反倾销刑事调查辩护策略，涉及潜在罚款逾亿元人民币", "协调香港律师通过法定追偿程序为数字营销公司追回2,000万美元", "代表欧洲食品企业处理290万美元跨境销售合同纠纷", "代理国有建设集团处理近20起建设工程合同纠纷"], matters_en=["Developed a defense strategy for a U.S. trading company in an anti-dumping criminal investigation involving potential fines exceeding RMB 100 million", "Coordinated with Hong Kong counsel to recover USD 20 million for a digital-marketing company", "Represented a European food company in a USD 2.9 million cross-border sales dispute", "Represented a state-owned construction group in nearly 20 construction-contract disputes"],
         education_zh=["香港中文大学 法律博士（J.D.）（2020年）", "瑞士日内瓦大学 国际法与欧盟法硕士（LL.M.）（2017年）", "山东大学 法学学士、文学学士（英语）（2016年）"], education_en=["J.D., The Chinese University of Hong Kong (2020)", "LL.M. in International Law and EU Law, University of Geneva (2017)", "LL.B. in Law & B.A. in English, Shandong University (2016)"],
         bars_zh="中华人民共和国执业律师；美国加州执业律师；英格兰与威尔士事务律师", bars_en="PRC Bar; State Bar of California; Solicitor of the Senior Courts of England and Wales",
         recognition_zh=["广东省五一劳动奖章（2024年）", "深圳市跨境法律业务“新星律师”", "广东省法律职业技能竞赛一等奖"], recognition_en=["Guangdong Province May 1st Labor Medal (2024)", "Rising Star in Cross-Border Legal Practice, Shenzhen", "First Prize, Guangdong Legal Skills Competition"]),
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.72); sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)
    add_footer(sec)
    normal = doc.styles["Normal"]
    normal.font.name = CH_FONT; normal._element.rPr.rFonts.set(qn("w:eastAsia"), CH_FONT); normal.font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2"):
        s = doc.styles[style_name]; s.font.name = CH_FONT; s._element.rPr.rFonts.set(qn("w:eastAsia"), CH_FONT)
    # Cover: editorial profile cover with restrained brand palette.
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, before=30, after=36)
    p.add_run().add_picture(str(ROOT / "images/logo.png"), width=Inches(1.5))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, after=5)
    set_run(p.add_run("博衡律师事务所"), 27, True, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, after=24)
    set_run(p.add_run("CATHAY PRAXIS LAW, PC"), 19, True, GOLD, EN_FONT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, after=7)
    set_run(p.add_run("律所介绍与律师团队"), 16, False, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, after=50)
    set_run(p.add_run("FIRM PROFILE & PROFESSIONALS"), 11, False, GRAY, EN_FONT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; style_para(p, after=4)
    set_run(p.add_run("跨境有道 · 判断有衡"), 13, True, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("CROSS-BORDER CLARITY. PRACTICAL JUDGMENT."), 9.5, color=GRAY, font=EN_FONT)
    p.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "关于博衡", "ABOUT CATHAY PRAXIS LAW")
    add_text(doc, "博衡律师事务所（Cathay Praxis Law, PC）是一家立足美国、专注于跨境法律服务的律师事务所。我们协助企业、投资机构与企业家处理跨境投资并购、美国市场准入与监管合规、出口管制与经济制裁、国际争议解决、知识产权以及其他涉美法律事务。", "Cathay Praxis Law, PC is a U.S.-based law firm focused on cross-border legal services. We advise companies, investors, and entrepreneurs on cross-border investment and M&A, U.S. market entry and regulatory compliance, export controls and economic sanctions, international dispute resolution, intellectual property, and other U.S.-related legal matters.")
    add_text(doc, "本所团队成员拥有中国、美国及英格兰与威尔士等多法域执业背景，并曾在国际律师事务所、领先中资律师事务所及专业机构工作。依托对不同法律制度、监管环境与商业实践的理解，我们能够协调跨法域资源，帮助客户识别风险、优化交易架构、解决争议并推进商业目标。", "Our professionals are qualified across multiple jurisdictions, including China, the United States, and England & Wales, and have experience at international law firms, leading Chinese firms, and professional institutions. Drawing on an integrated understanding of legal systems, regulatory environments, and commercial practice, we coordinate across jurisdictions to help clients identify risk, structure transactions, resolve disputes, and advance business objectives.")
    add_text(doc, "我们相信，优秀的法律服务不仅在于准确理解规则，更在于把法律判断转化为清晰、及时且可执行的商业解决方案。", "We believe excellent legal service requires more than an accurate reading of the rules: it requires turning legal judgment into clear, responsive, and executable business solutions.", indent=False)

    add_heading(doc, "核心业务领域", "CORE PRACTICE AREAS")
    add_bilingual_bullets(doc,
        ["跨境交易及美国市场准入", "出口管制、经济制裁与国际贸易合规", "国际仲裁与跨境争议解决", "知识产权诉讼、商标申请及跨境知识产权战略", "公司治理、监管合规与风险管理"],
        ["Cross-Border Transactions & U.S. Market Entry", "Export Controls, Economic Sanctions & International Trade Compliance", "International Arbitration & Cross-Border Dispute Resolution", "IP Litigation, Trademark Prosecution & Cross-Border IP Strategy", "Corporate Governance, Regulatory Compliance & Risk Management"])
    add_heading(doc, "服务对象与行业", "CLIENTS & INDUSTRIES", 2)
    add_bilingual_bullets(doc,
        ["开展美国及全球业务的中国企业与上市公司", "跨国公司、投资机构、私募基金与企业家", "科技、先进制造、工业、能源与基础设施企业", "国际贸易、供应链、消费品牌及知识产权密集型企业"],
        ["Chinese companies and listed enterprises operating in the United States and globally", "Multinational companies, investors, private equity funds, and entrepreneurs", "Technology, advanced manufacturing, industrial, energy, and infrastructure businesses", "International trade, supply-chain, consumer-brand, and IP-intensive businesses"])

    add_heading(doc, "联系我们", "CONTACT", 2)
    t = doc.add_table(rows=4, cols=2); t.autofit = False
    widths = [1800, 7560]
    for row, (label, val) in zip(t.rows, [("邮寄地址 / Mailing Address", "1259 El Camino Real, Unit #1468, Menlo Park, CA 94025"), ("电话 / Phone", "+1 341 237 9530"), ("邮箱 / Email", "info@cathaypraxislaw.com"), ("网站 / Website", "www.cathaypraxislaw.com")]):
        for i, c in enumerate(row.cells): c.width = Inches(widths[i]/1440); set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(label), 9, True, NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(val), 9.3, color=GRAY, font=EN_FONT)

    for person in people:
        doc.add_page_break()
        add_profile_header(doc, person)
        add_heading(doc, "专业概述", "PROFILE", 2)
        add_text(doc, person["bio_zh"], person["bio_en"], indent=False)
        add_heading(doc, "主要业务领域", "PRACTICE AREAS", 2)
        add_bilingual_bullets(doc, person["practices_zh"], person["practices_en"])
        add_heading(doc, "代表经验", "SELECTED EXPERIENCE", 2)
        add_bilingual_bullets(doc, person["matters_zh"], person["matters_en"])
        add_heading(doc, "教育背景", "EDUCATION", 2)
        add_bilingual_bullets(doc, person["education_zh"], person["education_en"])
        add_heading(doc, "执业资格", "ADMISSIONS", 2)
        add_text(doc, person["bars_zh"], person["bars_en"], indent=False)
        add_heading(doc, "荣誉与经历", "RECOGNITION & EXPERIENCE", 2)
        add_bilingual_bullets(doc, person["recognition_zh"], person["recognition_en"])

    p = doc.add_paragraph(); style_para(p, before=18, after=0)
    set_run(p.add_run("注：本介绍依据博衡律师事务所网站资料整理，内容仅供律所及团队介绍之用，不构成法律意见。"), 8.2, color=GRAY)
    p = doc.add_paragraph(); style_para(p, after=0)
    set_run(p.add_run("Note: This profile is based on information published on the Cathay Praxis Law website. It is for general firm and professional background purposes only and does not constitute legal advice."), 7.8, color=GRAY, font=EN_FONT)
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
